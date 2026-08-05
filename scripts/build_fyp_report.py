# -*- coding: utf-8 -*-
"""Generate the completed KNUST FYP report for Tarrlok."""

from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"C:\laragon\www\bbbms-project")
OUT = ROOT / "Final_Year_Project_Report_Blockchain_based_blood_bank_management.docx"
FIG_DIR = ROOT / "scripts" / "report_figures"
FIG_DIR.mkdir(parents=True, exist_ok=True)


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, size=12, bold=False, italic=False, align="left", space_after=8, first_line=True, space_before=0):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line and align == "justify":
        p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=11)
    doc.add_paragraph()


def caption(doc, text):
    add_para(doc, text, bold=True, align="center", first_line=False, space_before=6, space_after=10)


def add_picture(doc, path, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))


def add_page_number(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=10)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def box(ax, x, y, w, h, text, fc="#F4EFE6", ec="#1B2A4A"):
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.08", linewidth=1.4, facecolor=fc, edgecolor=ec)
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=8.5, color="#1B2A4A", wrap=True)


def arrow(ax, x1, y1, x2, y2):
    ax.add_patch(FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=12, linewidth=1.2, color="#8B1E2D"))


def make_figures():
    plt.rcParams["font.family"] = "DejaVu Sans"

    fig, ax = plt.subplots(figsize=(9.2, 5.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.2)
    ax.axis("off")
    ax.set_title("Four-layer architecture of Tarrlok", fontsize=12, pad=8)
    box(ax, 0.4, 5.0, 9.2, 0.9, "Presentation  |  Blade + CSS  |  Landing, portals, /track, shared ledger", fc="#F8EDE3")
    box(ax, 0.4, 3.6, 9.2, 0.9, "Application  |  Laravel 13 controllers, middleware, notifications, Blockchain* services", fc="#E8EEF6")
    box(ax, 0.4, 1.7, 4.4, 1.5, "Data\nMySQL 8\nusers, hospitals, donors,\nunits, requests, tamper attempts", fc="#EAF4EA")
    box(ax, 5.2, 1.7, 4.4, 1.5, "Blockchain\nHardhat + BloodBank.sol\nregister / screen / issue\ngetUnit + event log", fc="#F8E6E8")
    box(ax, 0.4, 0.25, 9.2, 1.1, "Node bridge: anchor-event.js  |  read-ledger.js  |  chain-status.js   (ethers.js v6)", fc="#FFF8E7")
    arrow(ax, 5, 5.0, 5, 4.5)
    arrow(ax, 5, 3.6, 5, 3.2)
    arrow(ax, 2.6, 1.7, 2.6, 1.35)
    arrow(ax, 7.4, 1.7, 7.4, 1.35)
    p = FIG_DIR / "fig3_1_architecture.png"
    fig.tight_layout()
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(9.4, 3.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 3.2)
    ax.axis("off")
    ax.set_title("Blood unit lifecycle", fontsize=12, pad=6)
    steps = [
        (0.2, "Lab registers\nunit + slip"),
        (2.5, "Quarantine\nscreening pending"),
        (4.8, "Screen HIV/HBV\nHCV/syphilis"),
        (7.1, "Cleared stock\nor discarded"),
        (9.4, "Partner issue\nFIFO + on-chain"),
    ]
    for x, label in steps:
        box(ax, x, 1.15, 2.1, 1.35, label, fc="#F4EFE6")
    for x in (2.3, 4.6, 6.9, 9.2):
        arrow(ax, x, 1.82, x + 0.2, 1.82)
    ax.text(6, 0.35, "Each successful transition is anchored. Failed contract calls are stored as blocked attempts.", ha="center", fontsize=8, color="#333333")
    p = FIG_DIR / "fig3_2_lifecycle.png"
    fig.tight_layout()
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(9.2, 2.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 2.4)
    ax.axis("off")
    ax.set_title("Write path: Laravel to BloodBank.sol", fontsize=12, pad=6)
    labels = ["Controller\nMySQL write", "BlockchainService", "anchor-event.js", "Hardhat :8545", "BloodBank.sol\ntx hash saved"]
    for i, label in enumerate(labels):
        box(ax, 0.25 + i * 2.35, 0.55, 2.15, 1.35, label, fc="#E8EEF6" if i < 2 else "#F8E6E8")
        if i < 4:
            arrow(ax, 2.4 + i * 2.35, 1.2, 2.55 + i * 2.35, 1.2)
    p = FIG_DIR / "fig4_1_write.png"
    fig.tight_layout()
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(9.2, 2.8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 2.4)
    ax.axis("off")
    ax.set_title("Shared-ledger read path", fontsize=12, pad=6)
    labels = ["Admin / hospital\n/ lab portal", "LedgerService", "read-ledger.js", "eth_getLogs\n+ getUnit()", "Same Blade\nledger UI"]
    for i, label in enumerate(labels):
        box(ax, 0.25 + i * 2.35, 0.55, 2.15, 1.35, label, fc="#EAF4EA")
        if i < 4:
            arrow(ax, 2.4 + i * 2.35, 1.2, 2.55 + i * 2.35, 1.2)
    p = FIG_DIR / "fig4_2_read.png"
    fig.tight_layout()
    fig.savefig(p, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    return {
        "arch": FIG_DIR / "fig3_1_architecture.png",
        "life": FIG_DIR / "fig3_2_lifecycle.png",
        "write": FIG_DIR / "fig4_1_write.png",
        "read": FIG_DIR / "fig4_2_read.png",
    }


def build():
    figs = make_figures()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Tarrlok FYP  |  KNUST Computer Engineering  |  Page ")
    set_run_font(run, size=10)
    add_page_number(fp)

    add_para(doc, "KWAME NKRUMAH UNIVERSITY OF SCIENCE AND TECHNOLOGY, KUMASI", size=14, bold=True, align="center", first_line=False, space_after=4)
    add_para(doc, "COLLEGE OF ENGINEERING", size=13, bold=True, align="center", first_line=False, space_after=4)
    add_para(doc, "DEPARTMENT OF COMPUTER ENGINEERING", size=13, bold=True, align="center", first_line=False, space_after=28)
    add_para(doc, "DESIGNING AND DEVELOPING A BLOCKCHAIN-BASED BLOOD BANK MANAGEMENT SYSTEM", size=16, bold=True, align="center", first_line=False, space_after=28)
    add_para(doc, "By", size=12, align="center", first_line=False, space_after=8)
    add_para(doc, "OFEI-PALM VALENTINO PAPA AYITEY  (1825922)", size=13, bold=True, align="center", first_line=False, space_after=4)
    add_para(doc, "ASIEDU ENOCH OFORI KWASI  (1818522)", size=13, bold=True, align="center", first_line=False, space_after=18)
    add_para(doc, "A Project Report submitted to the Department of Computer Engineering, Kwame Nkrumah University of Science and Technology, in partial fulfilment of the requirements for the award of the degree of", size=12, align="center", first_line=False, space_after=12)
    add_para(doc, "BACHELOR OF SCIENCE (COMPUTER ENGINEERING)", size=13, bold=True, align="center", first_line=False, space_after=24)
    add_para(doc, "Supervisor", size=12, align="center", first_line=False, space_after=4)
    add_para(doc, "Ing. Dr Bright Yeboah-Akowuah", size=13, bold=True, align="center", first_line=False, space_after=24)
    add_para(doc, "August 2026", size=12, align="center", first_line=False, space_after=24)

    doc.add_page_break()
    heading(doc, "Declaration", 1)
    add_para(doc, "We hereby declare that this submission is our own work towards the BSc. Computer Engineering degree and that, to the best of our knowledge, it contains no material previously published by another person nor material which has been accepted for the award of any other degree of the University, except where due acknowledgement has been made in the text.", align="justify")
    add_para(doc, "Ofei-Palm Valentino Papa Ayitey  (1825922)  ______________________     Date: ______________", first_line=False, space_before=18)
    add_para(doc, "Asiedu Enoch Ofori Kwasi  (1818522)  ______________________     Date: ______________", first_line=False, space_before=12)
    add_para(doc, "Certified by:", first_line=False, space_before=18, bold=True)
    add_para(doc, "Ing. Dr Bright Yeboah-Akowuah (Supervisor)  ______________________     Date: ______________", first_line=False, space_before=8)

    doc.add_page_break()
    heading(doc, "Abstract", 1)
    add_para(doc, "Blood is a critical clinical resource. Globally, national blood systems still struggle with inequitable access, incomplete quality assurance and weak digital coordination [1], [2]. In Ghana, collection, screening, storage and distribution remain fragmented across more than one hundred and fifty hospital-based transfusion services alongside the National Blood Service [3], [4]. That fragmentation makes it difficult for facilities to see one another's stock, to prove that a unit was not altered after registration, and to let a donor follow a donation without exposing other people's records.", align="justify")
    add_para(doc, "This project designed and implemented Tarrlok, a blockchain-based blood bank management system for Health Facilities Regulatory Agency (HeFRA) licensed hospitals. Day-to-day operations run on Laravel and MySQL. Critical lifecycle events\u2014unit registration, laboratory screening and partner issue\u2014are anchored on a local Ethereum smart contract (BloodBank.sol) deployed with Hardhat. Each physical bag receives a unique unit code (for example UNIT-002-00001) that is hashed on-chain. Hospital, laboratory and platform-admin portals all read the same shared ledger, compare operational records with on-chain state, and list blocked write attempts. Donors track a single unit publicly by code, without creating an account.", align="justify")
    add_para(doc, "The system was demonstrated with seeded Korle Bu Teaching Hospital and Ridge Hospital accounts, automated PHPUnit feature tests, and an end-to-end workflow from registration through screening, partner transfer, public tracking and tamper detection. The work shows that a hybrid architecture\u2014mutable operational data off-chain and an immutable audit log on-chain\u2014is a practical way to improve traceability and multi-stakeholder visibility in a Ghanaian hospital network without requiring every facility to run its own blockchain node.", align="justify")
    add_para(doc, "Keywords: blockchain; blood bank management; Ethereum; smart contracts; Laravel; Ghana; HeFRA; traceability.", italic=True, first_line=False)

    doc.add_page_break()
    heading(doc, "Acknowledgements", 1)
    add_para(doc, "We first thank Almighty God for the strength and wisdom granted throughout this project and our studies at KNUST. We are grateful to our supervisor, Ing. Dr Bright Yeboah-Akowuah, for guidance, constructive criticism and encouragement from proposal to completion. We thank the faculty and staff of the Department of Computer Engineering for the teaching and resources that made the work possible. Finally, we thank our families for their love, sacrifice and belief in us. We dedicate this report to them.", align="justify")

    heading(doc, "Table of Contents", 1)
    toc = [
        "Declaration",
        "Abstract",
        "Acknowledgements",
        "List of Figures",
        "List of Tables",
        "List of Abbreviations",
        "Chapter 1  Introduction",
        "    1.1  Background of the Study",
        "    1.2  Problem Statement",
        "    1.3  Aim of the Study",
        "    1.4  Objectives of the Study",
        "    1.5  Significance of the Study",
        "    1.6  Scope of the Study",
        "    1.7  Organisation of the Report",
        "Chapter 2  Literature Review",
        "    2.1  Introduction",
        "    2.2  Blood Bank Management Systems",
        "    2.3  Blockchain Technology",
        "    2.4  Blockchain in Healthcare",
        "    2.5  Blockchain in Blood Bank Management",
        "    2.6  Research Gaps and Justification",
        "    2.7  Summary",
        "Chapter 3  Methodology and System Design",
        "    3.1  Software Development Approach",
        "    3.2  Requirements Analysis",
        "    3.3  Architectural Design",
        "    3.4  Data Model and Smart Contract Design",
        "    3.5  Security and Integrity Design",
        "Chapter 4  Implementation",
        "    4.1  Development Environment",
        "    4.2  Application Modules and Roles",
        "    4.3  Blood Unit Lifecycle",
        "    4.4  Blockchain Integration",
        "    4.5  Shared Ledger and Tamper Detection",
        "Chapter 5  Testing, Results and Discussion",
        "    5.1  Test Strategy",
        "    5.2  Functional Demonstration",
        "    5.3  Integrity and Blockchain Results",
        "    5.4  Discussion",
        "Chapter 6  Conclusion and Recommendations",
        "    6.1  Conclusion",
        "    6.2  Limitations",
        "    6.3  Recommendations for Future Work",
        "References",
        "Appendix A  Demo Accounts and Commands",
    ]
    for item in toc:
        add_para(doc, item, first_line=False, space_after=2)

    heading(doc, "List of Figures", 1)
    add_para(doc, "Figure 3.1  Four-layer architecture of Tarrlok.", first_line=False, space_after=2)
    add_para(doc, "Figure 3.2  Blood unit lifecycle from registration to partner issue or discard.", first_line=False, space_after=2)
    add_para(doc, "Figure 4.1  Write path from a Laravel controller to BloodBank.sol.", first_line=False, space_after=2)
    add_para(doc, "Figure 4.2  Shared-ledger read path used by admin, hospital and laboratory portals.", first_line=False, space_after=2)

    heading(doc, "List of Tables", 1)
    add_para(doc, "Table 2.1  Comparison of related blockchain blood-bank systems with Tarrlok.", first_line=False, space_after=2)
    add_para(doc, "Table 3.1  Functional requirements mapped to implemented modules.", first_line=False, space_after=2)
    add_para(doc, "Table 3.2  On-chain events versus MySQL fields.", first_line=False, space_after=2)
    add_para(doc, "Table 4.1  User roles and portal capabilities.", first_line=False, space_after=2)
    add_para(doc, "Table 5.1  Representative automated tests and outcomes.", first_line=False, space_after=2)

    heading(doc, "List of Abbreviations", 1)
    abbr = [
        ("AABB", "Association for the Advancement of Blood and Biotherapies"),
        ("BSIS", "Blood Safety Information System"),
        ("CSRF", "Cross-Site Request Forgery"),
        ("EVM", "Ethereum Virtual Machine"),
        ("FIFO", "First In, First Out"),
        ("HeFRA", "Health Facilities Regulatory Agency"),
        ("HBV / HCV", "Hepatitis B / Hepatitis C virus"),
        ("HIV", "Human Immunodeficiency Virus"),
        ("KNUST", "Kwame Nkrumah University of Science and Technology"),
        ("NBS", "National Blood Service, Ghana"),
        ("RPC", "Remote Procedure Call"),
        ("TTI", "Transfusion-Transmissible Infection"),
        ("WHO", "World Health Organization"),
    ]
    for k, v in abbr:
        add_para(doc, f"{k}  \u2014  {v}", first_line=False, space_after=2)

    # CHAPTER 1
    doc.add_page_break()
    heading(doc, "Chapter 1  Introduction", 1)
    heading(doc, "1.1  Background of the Study", 2)
    add_para(doc, "Blood banks collect, screen, store and issue blood for transfusion. They therefore need reliable information systems that preserve the identity of each unit, screening results and movement between facilities [5]. The World Health Organization continues to emphasise national systems that combine voluntary unpaid donation, quality-assured screening of every unit for HIV, hepatitis B, hepatitis C and syphilis, and accountable distribution [1], [2]. WHO's June 2026 fact sheet and the Global status report on blood safety and availability 2025 still document large gaps in inspection, licensing and digital coordination of blood services, especially in low- and middle-income settings [1], [2].", align="justify")
    add_para(doc, "In Ghana, the National Blood Service coordinates policy and regional centres under the National Blood Policy and the National Blood Service Act, 2022 (Act 1042) [3], [6]. At the same time, more than one hundred and fifty hospitals still operate their own hospital-based transfusion services [3], [4]. Dei and Ansah describe the practical difficulty of coordinating quality and adequacy across this mixed national and hospital-based model [4]. Health Facilities Regulatory Agency (HeFRA) licensing under the Health Institutions and Facilities Act, 2011 (Act 829) already provides a legal identity for facilities that can join a shared digital network [7], [8].", align="justify")
    add_para(doc, "Computerised tools such as the Blood Safety Information System improved on paper logs, yet hospital systems remain largely centralised. A change made in one database is not automatically visible, or verifiable, at another hospital. Blockchain is a cryptographically linked ledger in which confirmed records are expensive to rewrite [9], [10], [11]. Ethereum extended that model with smart contracts\u2014programs that execute deterministic rules when called [12], [13]. For blood banking, the useful property is not cryptocurrency, but a shared, time-stamped log of who registered, screened or issued a unit. Sensitive donor details need not sit on-chain; hashes and lifecycle facts can, while MySQL remains the operational store [14], [15].", align="justify")
    add_para(doc, "This study therefore built Tarrlok: a Ghana-oriented hospital network in which HeFRA-licensed facilities register, laboratory staff record and screen units, hospitals exchange stock, donors track one unit by public code, and every stakeholder portal can read the same on-chain audit log.", align="justify")

    heading(doc, "1.2  Problem Statement", 2)
    add_para(doc, "Existing blood bank records in many Ghanaian facilities are centralised or disconnected. That produces four practical problems. First, traceability breaks when a unit leaves the registering laboratory: partner hospitals cannot independently verify screening or ownership. Second, operational databases can be edited; there is no automatic comparison against an immutable copy of the original blood group, expiry or screening outcome. Third, coordination is weak, so one hospital may discard near-expiry stock while another searches for the same group. Fourth, donors have little honest visibility: either they are shut out entirely, or a donor portal pretends they have an account when the real workflow is a slip with a unit identifier.", align="justify")
    add_para(doc, "A system is needed that keeps hospital operations usable (login, inventory, requests, screening) while giving all network stakeholders a common, tamper-evident view of critical events and of attempts that the ledger rejected.", align="justify")

    heading(doc, "1.3  Aim of the Study", 2)
    add_para(doc, "The aim of this study is to design and develop a blockchain-based blood bank management system that improves traceability, transparency, security and coordination across HeFRA-licensed hospital blood banks in Ghana.", align="justify")

    heading(doc, "1.4  Objectives of the Study", 2)
    add_para(doc, "The specific objectives were to:", first_line=False)
    add_bullets(doc, [
        "Design a hybrid architecture in which MySQL supports operations and an Ethereum smart contract stores an immutable audit of unit registration, screening and issue.",
        "Implement secure facility onboarding (HeFRA licence review), role-based hospital and laboratory portals, partner blood requests, expiry handling and public unit tracking by unique code.",
        "Anchor each physical unit identifier on-chain, enforce lifecycle guards in Solidity, and expose the same ledger, integrity alerts and blocked attempts to admin, hospital and laboratory stakeholders.",
        "Evaluate the system through automated tests and a multi-hospital demonstration covering happy-path transfer and tamper detection.",
    ])

    heading(doc, "1.5  Significance of the Study", 2)
    add_para(doc, "Patients benefit when issued units can be shown to have passed screening and not to have been silently rewritten after registration. Donors benefit from a simple, honest tracking path that uses the code on the donation slip rather than a fake login. Hospital and laboratory staff gain a working inventory and request workflow plus a shared ledger instead of an admin-only blockchain screen. Platform administrators and, by extension, regulators gain an audit trail that is harder to rewrite than a standalone MySQL row. Academically, the project documents a permissioned consortium pattern that is feasible for a final-year deployment: one local chain, many application nodes (stakeholder portals) reading the same log.", align="justify")

    heading(doc, "1.6  Scope of the Study", 2)
    add_para(doc, "The implemented system covers hospital registration and HeFRA-oriented review, laboratory unit registration with donor phone linkage, four-marker screening (HIV, hepatitis B, hepatitis C, syphilis), 35-day whole-blood shelf life and expiry jobs, partner requests (approve, reject, issue, cancel), staff trace, public /track, blockchain anchoring, shared ledger visibility and integrity comparison. It does not perform clinical cross-match or the transfusion procedure itself, does not store full electronic health records, does not run a public Ethereum mainnet or per-hospital Geth nodes, and does not claim cryptographic proof of a direct SQL edit\u2014only that the operational row no longer matches the chain.", align="justify")

    heading(doc, "1.7  Organisation of the Report", 2)
    add_para(doc, "Chapter 2 reviews blood bank systems, blockchain fundamentals and related work. Chapter 3 presents methodology and design. Chapter 4 describes implementation. Chapter 5 reports testing and discussion. Chapter 6 concludes and recommends future work.", align="justify")

    # CHAPTER 2
    heading(doc, "Chapter 2  Literature Review", 1)
    heading(doc, "2.1  Introduction", 2)
    add_para(doc, "This chapter reviews traditional and computerised blood bank systems, blockchain and smart contracts, healthcare applications, and prior blockchain blood-bank proposals. Limitations of those works motivate Tarrlok.", align="justify")

    heading(doc, "2.2  Blood Bank Management Systems", 2)
    heading(doc, "2.2.1  Traditional and manual systems", 3)
    add_para(doc, "Historically, blood banks tracked donors, inventory and transfusions in paper registers. Manual logs delay retrieval, make rare groups hard to locate in emergencies and leave personal data poorly protected [1], [4]. Human identification error has long been recognised as a major cause of ABO-incompatible transfusion [16], [17]. Practices also varied between facilities, so one hospital's register could not be trusted as a network-wide source of truth.", align="justify")

    heading(doc, "2.2.2  Computerised blood bank systems", 3)
    add_para(doc, "Computerised systems introduced databases, barcodes and electronic cross-match support, reducing some clerical error [17]. Ghana's move toward BSIS and national policy coordination is part of that computerisation [3], [4]. The remaining structural weakness is centralisation: one hospital database is authoritative only for itself. When a unit is transferred, the receiving site often cannot verify the original screening record independently. Unauthorised or accidental alteration of a central store is difficult for a partner facility to detect [18].", align="justify")

    heading(doc, "2.2.3  Limitations of existing operational systems", 3)
    add_para(doc, "Eder and Chambers note that incomplete traceability after a unit leaves the blood bank complicates investigation of reactions and contaminated products [18]. WHO guidance still stresses quality-assured screening and accountable national supply [1], [2]. In Ghana, hospital-based services that operate with limited integration into the National Blood Service make real-time visibility of neighbouring stock especially weak [3], [4]. These operational facts justify a digital network among licensed hospitals, not only a better spreadsheet inside one laboratory.", align="justify")

    heading(doc, "2.3  Blockchain Technology", 2)
    heading(doc, "2.3.1  Fundamentals", 3)
    add_para(doc, "A blockchain is a distributed ledger of ordered blocks, each containing transactions and a cryptographic hash of the previous block [9], [10], [11], [19]. Altering a confirmed block requires rewriting subsequent hashes and obtaining network acceptance, which is what gives the structure its tamper resistance. Crosby et al. summarised this beyond-Bitcoin value as a shared, append-only record among parties that do not fully trust one another [19]. Consensus algorithms such as proof-of-work or proof-of-stake remove the need for a single trusted operator on public networks [10], [20]. For enterprise and campus prototypes, a permissioned or locally simulated chain can still provide an append-only log that multiple application roles can query [21].", align="justify")

    heading(doc, "2.3.2  Smart contracts", 3)
    add_para(doc, "Nick Szabo described computerised transaction protocols that execute contract terms [22], [23]. Ethereum made general-purpose smart contracts practical through the Ethereum Virtual Machine and languages such as Solidity [12], [13]. A contract is addressed on-chain, stores state, exposes functions and emits events. Calls are transactions: they are ordered, timestamped and auditable. In a blood-bank setting, a contract can refuse to register the same unit twice, refuse a second screening, or refuse issue of an expired or uncleared unit\u2014rules that a lone MySQL UPDATE cannot enforce globally.", align="justify")

    heading(doc, "2.3.3  Types of blockchain", 3)
    add_para(doc, "Public chains (Bitcoin, public Ethereum) allow anyone to participate and maximise censorship resistance at the cost of throughput and fee volatility [9], [20]. Private chains are operated by one organisation and are faster, but they re-centralise governance [21]. Consortium chains are governed by several organisations\u2014an apt model for hospitals plus a regulator [24]. Hybrid designs keep some data private and some public. Tarrlok follows a consortium application model on a local Ethereum-compatible chain: stakeholders are authenticated hospital, laboratory and admin users, not anonymous mainnet wallets.", align="justify")

    heading(doc, "2.3.4  Limitations of blockchain", 3)
    add_para(doc, "Public networks process far fewer transactions per second than conventional payment switches [20]. Deployment cost and skills are non-trivial in low-resource health systems [14]. Permanence also conflicts with storing personally identifiable health data on-chain [25]. The accepted mitigation, which this project adopts, is to keep clinical and personal detail off-chain and place only verification records, unit codes, status and actor identifiers on-chain [14], [25].", align="justify")

    heading(doc, "2.4  Blockchain in Healthcare", 2)
    add_para(doc, "Healthcare surveys describe blockchain as a candidate for auditability, multi-party data sharing and reduced reliance on a single institutional database [15], [26]. Gordon and Catalini discuss patient-driven interoperability and the institutional barriers to data movement [27]. Yue et al. proposed keeping sensitive records off-chain while using the chain for access control or verification [25]. These ideas transfer directly to blood units: the bag is a physical asset whose digital twin should not be silently rewritten.", align="justify")
    add_para(doc, "Chakraborty, Aich and Kim presented a blockchain-oriented healthcare framework focused on trusted data exchange among stakeholders at ICACT 2019 [28]. Their work is not a blood-bank product, but it supports the claim that healthcare multi-party workflows benefit from an append-only transaction layer. Blood itself is also a perishable supply-chain asset. Kshetri and Saberi et al. show why blockchain is attractive for provenance, multi-party visibility and reduced information asymmetry in supply chains\u2014properties that map cleanly onto hospital-to-hospital blood transfer [29], [30].", align="justify")

    heading(doc, "2.5  Blockchain in Blood Bank Management", 2)
    add_para(doc, "Shreshtha, Rajput and Singh argued for blockchain in blood-bank supply management to improve trust along the donation-to-transfusion path [31]. Pradhan, Singh and Kumar described a blockchain-enabled traceable transportation system for blood banks, emphasising movement visibility [32]. Sanga et al. proposed a smart blood-bank design on blockchain for real-time tracking and fraud prevention, including identity and incentive ideas beyond a basic inventory [33]. These studies show academic consensus that immutability and shared visibility matter, but most remain conceptual or isolated prototypes. Few describe a complete hospital information workflow (facility licensing, laboratory screening panels, partner requests, expiry jobs and a donor-facing lookup) together with an honest statement that the chain is an audit log rather than a full replacement for the hospital database.", align="justify")

    heading(doc, "2.6  Research Gaps and Justification", 2)
    add_para(doc, "Three gaps remain. First, end-to-end operational software for Ghanaian multi-hospital exchange is rarely combined with on-chain guards and a shared reader UI. Second, many prototypes either put too much personal data on-chain or hide the chain inside a single admin console, which recreates information asymmetry. Third, failed or illegal transitions are often only console errors; stakeholders never see who attempted a double screening or an issue of an expired unit. Tarrlok addresses these gaps with a working Laravel network, BloodBank.sol lifecycle rules, a shared ledger on every portal, integrity comparison of MySQL against getUnit(), and a tamper-attempt log keyed to the authenticated user.", align="justify")

    caption(doc, "Table 2.1  Comparison of related systems with the implemented project")
    add_table(
        doc,
        ["Work", "Platform", "Strengths", "Limitations relative to this project"],
        [
            ["Shreshtha et al. [31]", "Conceptual / blockchain supply", "Argues traceability need", "Not a full multi-role hospital product"],
            ["Pradhan et al. [32]", "Blockchain transport focus", "Movement visibility", "Limited hospital ops and screening workflow"],
            ["Sanga et al. [33]", "Blockchain tracking / fraud themes", "Real-time tracking narrative", "Broader identity/AI claims; not Ghana HeFRA ops"],
            ["Chakraborty et al. [28]", "Healthcare blockchain framework", "Multi-stakeholder trust model", "Not blood-unit lifecycle specific"],
            ["Tarrlok (this work)", "Laravel, MySQL, Solidity, Hardhat", "HeFRA onboarding, screening, exchange, shared ledger, public unit track", "Local chain; no mainnet or per-hospital node"],
        ],
    )

    heading(doc, "2.7  Summary", 2)
    add_para(doc, "Literature supports computerisation of blood banks and the use of blockchain as a tamper-evident multi-party log, provided personal data stay off-chain. Prior blood-bank blockchain papers motivate the problem but leave a gap for a deployable Ghana-oriented system in which every stakeholder can see the same ledger. Chapter 3 turns that gap into design choices.", align="justify")

    # CHAPTER 3
    heading(doc, "Chapter 3  Methodology and System Design", 1)
    heading(doc, "3.1  Software Development Approach", 2)
    add_para(doc, "The project used an Agile, sprint-based approach. Requirements evolved as hospital, laboratory and blockchain concerns were integrated: first authentication and facility registration, then unit and screening workflows, then partner exchange, then on-chain anchoring, then shared ledger visibility and integrity checks. Each sprint ended with runnable software rather than a paper-only design. Automated PHPUnit feature tests and manual demonstration scripts were used as regression gates [34].", align="justify")

    heading(doc, "3.2  Requirements Analysis", 2)
    add_para(doc, "Functional and non-functional requirements were derived from the literature, Ghana's facility-licensing context [7], [8], national blood policy [3], [6], and the operational realities of hospital blood banks [4]. WHO's four-marker screening recommendation (HIV, HBV, HCV, syphilis) was adopted as the laboratory panel [1]. Whole-blood shelf life was set at 35 days, matching common hospital practice encoded in config/tarrlok.php, with a seven-day expiry warning and a 56-day minimum donation interval.", align="justify")
    caption(doc, "Table 3.1  Functional requirements and implementation")
    add_table(
        doc,
        ["Requirement", "Implementation"],
        [
            ["Facility registration and approval", "Three-step wizard; admin approve/reject with email notices"],
            ["Role-based access", "admin, hospital, lab middleware and portals"],
            ["Donor linkage without donor login", "Donor record by phone; public /track by unit code"],
            ["Unit registration and unique ID", "UNIT-{hospitalId}-{sequence}; donation slip"],
            ["Screening", "HIV, HBV, HCV, syphilis; clear or fail; quarantine until clear"],
            ["Inventory and expiry", "35-day shelf life; hourly/manual expiry command"],
            ["Partner requests", "Create, approve, reject, issue (FIFO), cancel"],
            ["Blockchain audit", "registerUnit, recordScreening, recordIssue + tx hashes"],
            ["Shared visibility", "Ledger page on admin, hospital and lab portals"],
            ["Tamper evidence", "getUnit compare; blockchain_tamper_attempts"],
        ],
    )
    add_para(doc, "Non-functional requirements included confidentiality of donor contact data, integrity of lifecycle facts, availability of the operational app even if the chain is offline, auditability of critical writes, and a user interface simple enough for a viva demonstration.", align="justify")

    heading(doc, "3.3  Architectural Design", 2)
    add_para(doc, "Tarrlok uses a four-layer architecture, shown in Figure 3.1.", align="justify")
    add_picture(doc, figs["arch"])
    caption(doc, "Figure 3.1  Four-layer architecture of Tarrlok.")
    add_para(doc, "Presentation layer. Server-rendered Blade templates and plain CSS provide the landing page, login and registration, admin console, hospital portal, laboratory portal, staff trace and public track pages. There is no separate React single-page application. This choice reduced build complexity and kept authentication in ordinary server sessions [34].", align="justify")
    add_para(doc, "Application layer. Laravel 13 (PHP 8.3+) implements routing, validation, Eloquent models, notifications, scheduled expiry and blockchain services [34]. Controllers are grouped by role. BlockchainService spawns Node scripts that call the contract through ethers.js [35], [36].", align="justify")
    add_para(doc, "Data layer. MySQL stores hospitals, users, donors, blood units, requests, pivot issue rows and tamper attempts [37]. ACID transactions protect issue-time stock reservation. Transaction hashes live on blood_units so the operational UI can display them even when the node is down.", align="justify")
    add_para(doc, "Blockchain layer. A Hardhat local Ethereum node hosts BloodBank.sol [12], [13], [35]. The contract owner is the deployer wallet used by the application. Events UnitRegistered, UnitScreened and UnitIssued form the shared log. A view function getUnit(unitCode) returns existence, hospital id, blood group, expiry and screening enum.", align="justify")

    heading(doc, "3.4  Data Model and Smart Contract Design", 2)
    add_para(doc, "Core entities are Hospital (pending/approved/rejected, HeFRA licence id, region and institution type), User (admin/hospital/lab), Donor (phone-normalised, eligibility, tracking consent), BloodUnit (group, status, screening fields, expiry, three tx-hash columns) and BloodRequest (quantity, group, status, requesting and fulfilling hospitals).", align="justify")
    add_para(doc, "On-chain, each unit is keyed by keccak256(unitCode). UnitRecord stores exists, hospitalId, bloodGroup, expiresAt and screening. Guards reject duplicate registration, screening when not pending, unknown units, issue when not cleared, issue after expiry, wrong from-hospital, and empty actor names. Reverts cannot emit events; therefore blocked attempts are persisted in MySQL with the authenticated actor, while successful actions appear both as events and as stored hashes.", align="justify")
    caption(doc, "Table 3.2  On-chain versus off-chain data")
    add_table(
        doc,
        ["On-chain (BloodBank.sol)", "Off-chain (MySQL)"],
        [
            ["Unit code hash, group, expiry, screening enum, hospital id", "Donor name, phone, email, consent"],
            ["Actor id and name on events", "Full user accounts, roles, passwords"],
            ["Issue from/to hospital ids and request code", "Request urgency, notes, rejection reasons"],
            ["Tx hash implied by the mined transaction", "blockchain_*_tx columns for UI"],
        ],
    )
    add_picture(doc, figs["life"], width=6.3)
    caption(doc, "Figure 3.2  Blood unit lifecycle from registration to partner issue or discard.")

    heading(doc, "3.5  Security and Integrity Design", 2)
    add_para(doc, "Authentication uses hashed passwords and session cookies. Authorisation is enforced by role middleware and hospital scoping on laboratory writes. CSRF tokens protect state-changing forms [34]. Public track aborts unless the linked donor granted tracking consent, and only that unit is shown. Blockchain private keys stay on the server; stakeholders do not handle wallets. Integrity verification compares MySQL blood group, hospital id, screening and expiry with getUnit(). Mismatches are labelled tampered, with the last operational editor named as a best-effort clue, not as cryptographic attribution of a raw SQL update.", align="justify")
    add_para(doc, "Lifecycle in brief: the laboratory registers a unit (quarantine, screening pending, expires_at = collected_at + 35 days) and the chain emits UnitRegistered. Screening produces cleared/available or failed/discarded and UnitScreened. A hospital approves a partner request and issues FIFO cleared stock; unit hospital_id transfers and UnitIssued is recorded. An hourly job discards expired available units. Staff trace and public track read the timeline; all portals read the ledger.", align="justify")

    # CHAPTER 4
    heading(doc, "Chapter 4  Implementation", 1)
    heading(doc, "4.1  Development Environment", 2)
    add_para(doc, "The application was implemented as a monorepo: tarrlok/ for Laravel and blockchain/ for Hardhat. PHP 8.3+, Composer, MySQL 8 (or SQLite in automated tests), Node.js 18+ and Hardhat 2 with ethers v6 were used [34], [35], [36], [37]. Optional Docker Compose and Cloudflare Tunnel documentation support demonstration beyond localhost without exposing port 8545 publicly. Configuration lives in config/tarrlok.php (blood groups, screening markers, shelf life, Ghana regions, institution types) and config/blockchain.php (enable flag, RPC URL, private key, script paths).", align="justify")

    heading(doc, "4.2  Application Modules and Roles", 2)
    caption(doc, "Table 4.1  Roles")
    add_table(
        doc,
        ["Role", "Capabilities"],
        [
            ["Platform admin", "Approve/reject hospitals; blockchain health map; shared ledger; trace"],
            ["Hospital administrator", "Inventory, partners, requests, lab-staff accounts, facility profile, trace, ledger"],
            ["Laboratory staff", "Register units, screening reports, slips, inventory, trace, ledger"],
            ["Public donor", "No login; /track by unit code if consent is recorded"],
        ],
    )
    add_para(doc, "The public landing page presents live stock bars and impact counts from MySQL and routes guests to track, login or facility registration. Authenticated users are redirected to their role dashboard. Hospital registration is a three-step wizard capturing facility identity, HeFRA licence and administrator credentials. Approval and rejection send letter-style notifications (HospitalRegistrationApproved / Rejected). Donation status changes can notify a consenting donor by email when SMTP is configured; local development uses the log mailer.", align="justify")
    add_para(doc, "DemoSeeder creates two approved Greater Accra facilities\u2014Korle Bu Teaching Hospital (HFRA-GAR-2024) and Ridge Hospital (HFRA-GAR-2025)\u2014with hospital and laboratory users, sample donors and seeded units including the public demo code UNIT-002-00001.", align="justify")

    heading(doc, "4.3  Blood Unit Lifecycle", 2)
    add_para(doc, "Laboratory staff look up or create a donor by phone, then register a unit. The code format UNIT-{padded hospital id}-{sequence} is printed on a slip for the donor. Screening requires an explicit clear or fail action; clearing demands all four markers non-reactive. Only cleared, non-expired units count as issuable stock. Partner hospitals browse approved facilities, submit requests, and fulfilling hospitals approve and issue. Issue runs inside a database transaction with row locks to prevent double allocation, then attempts an on-chain recordIssue. If the chain rejects or is offline, inventory transfer can still complete and a warning plus tamper-attempt or missing-hash gap is recorded, which is honest about operational continuity.", align="justify")
    add_para(doc, "Expiry is handled by php artisan blood:mark-expired and by the Laravel scheduler (hourly). Units past shelf life that are still marked available are discarded so they cannot be issued. Near-expiry warnings use the seven-day window from configuration.", align="justify")

    heading(doc, "4.4  Blockchain Integration", 2)
    add_picture(doc, figs["write"])
    caption(doc, "Figure 4.1  Write path from a Laravel controller to BloodBank.sol.")
    add_para(doc, "After a successful business write, BlockchainService::anchor() JSON-encodes the action and arguments, runs node blockchain/scripts/anchor-event.js, and stores the returned 0x transaction hash on the unit (blockchain_register_tx, blockchain_screening_tx or blockchain_issue_tx). Configuration is environment-driven (BLOCKCHAIN_ENABLED, RPC URL, deployer private key). The first Hardhat development account is used locally and must never be reused in production [35]. Restarting the Hardhat node wipes chain state; the contract must be redeployed. Old MySQL hashes then fail integrity checks, which is itself a demonstration of mismatch detection.", align="justify")

    heading(doc, "4.5  Shared Ledger and Tamper Detection", 2)
    add_picture(doc, figs["read"])
    caption(doc, "Figure 4.2  Shared-ledger read path used by admin, hospital and laboratory portals.")
    add_para(doc, "BlockchainLedgerService calls read-ledger.js, which uses queryFilter on UnitRegistered, UnitScreened and UnitIssued and optionally getUnit for selected codes. Admin, hospital and laboratory routes all render the same Blade partials (shared/blockchain/index.blade.php and ledger.blade.php). Integrity alerts list database-versus-chain mismatches. Blocked attempts list actor name, role, hospital, action, unit code and revert reason from the blockchain_tamper_attempts table.", align="justify")
    add_para(doc, "This design answers the project requirement that all nodes see what is going on. In Tarrlok, nodes are authenticated stakeholder portals reading one permissioned chain, not separate proof-of-work miners at every hospital. That is an intentional, defensible consortium interpretation for a campus and hospital-IT deployment [21], [24].", align="justify")

    # CHAPTER 5
    heading(doc, "Chapter 5  Testing, Results and Discussion", 1)
    heading(doc, "5.1  Test Strategy", 2)
    add_para(doc, "Testing combined automated feature tests (PHPUnit with an in-memory SQLite database) and a scripted manual demonstration. Automated tests cover landing availability, authentication, hospital registration steps, shared-ledger authorisation, failed-anchor persistence, integrity comparison and navigation destinations [34]. BlockchainProcess calls are faked in PHPUnit so tests do not require a live Hardhat node. Manual tests exercise seeded Korle Bu and Ridge accounts against a running node.", align="justify")
    caption(doc, "Table 5.1  Representative automated tests")
    add_table(
        doc,
        ["Test focus", "Result"],
        [
            ["Guest landing, login and public /track render", "Pass"],
            ["Login redirects by role; invalid password rejected", "Pass"],
            ["Hospital registration wizard step order enforced", "Pass"],
            ["Guest denied ledger routes", "Pass"],
            ["Hospital, lab and admin ledger pages render", "Pass"],
            ["Failed chain write inserts tamper-attempt row visible to other roles", "Pass"],
            ["Integrity compare flags blood-group mismatch as tampered", "Pass"],
            ["Landing and dashboards expose live named routes", "Pass"],
        ],
    )

    heading(doc, "5.2  Functional Demonstration", 2)
    add_para(doc, "A standard viva script is: confirm chain health as platform admin; inspect Ridge cleared inventory; as Korle Bu request O+ from Ridge; as Ridge approve and issue; confirm Korle Bu inventory received the unit; open public /track with UNIT-002-00001; open staff Trace on the same code; open Network ledger on hospital and lab accounts and observe the same events. Laboratory registration of a new unit followed by screening produces fresh UnitRegistered and UnitScreened rows visible to every portal.", align="justify")

    heading(doc, "5.3  Integrity and Blockchain Results", 2)
    add_para(doc, "When the node is healthy, registration, screening and issue produce 0x-prefixed transaction hashes and decoded events with actor names. A second screening attempt is rejected by the contract and appears under Blocked attempts with the laboratory user's name. Editing a unit's blood group in MySQL after anchoring produces a Tampered integrity alert on all portals while the chain still reports the original group. That is the intended evidence story: the operational database is mutable; the ledger is not.", align="justify")

    heading(doc, "5.4  Discussion", 2)
    add_para(doc, "The hybrid model proved more realistic than putting the whole blood bank on blockchain. Hospitals still need rich queries, forms and eligibility rules that do not belong in Solidity. Conversely, relying on MySQL alone would not let a partner facility detect silent edits. Shared ledger UI was essential: without it, the chain would have remained an admin novelty and would have contradicted the multi-node visibility objective.", align="justify")
    add_para(doc, "Limitations observed in testing are consistent with the literature. A local Hardhat chain is not a national production network [20], [35]. Best-effort actor strings on events are only as trustworthy as the application that submitted them, because a single owner wallet signs all transactions. Continuity during chain outage means some operational writes lack anchors until replay is designed. These are documented trade-offs, not hidden defects. Relative to related blood-bank blockchain papers [31]\u2013[33], Tarrlok's contribution is a complete Ghana-oriented operational product plus an honest consortium ledger, not a claim of mainnet decentralisation or biometric/AI extras that were not built.", align="justify")

    # CHAPTER 6
    heading(doc, "Chapter 6  Conclusion and Recommendations", 1)
    heading(doc, "6.1  Conclusion", 2)
    add_para(doc, "This project designed and implemented Tarrlok, a blockchain-based blood bank management system tailored to HeFRA-licensed Ghanaian hospitals. The aim and objectives were met: a hybrid Laravel\u2013MySQL\u2013Ethereum architecture is running; facility onboarding, laboratory screening, partner exchange, expiry handling and public unit tracking are implemented; unique unit codes are anchored on-chain with lifecycle guards; and every stakeholder portal can inspect the same ledger, integrity alerts and blocked attempts. The work demonstrates that tamper-evident multi-hospital blood tracking is achievable without discarding conventional hospital information-system practice.", align="justify")

    heading(doc, "6.2  Limitations", 2)
    add_bullets(doc, [
        "The blockchain is a local Hardhat network signed by one application wallet, not a production consortium with per-hospital keys.",
        "Email delivery depends on SMTP configuration; local development uses log mail.",
        "No clinical cross-match, cold-chain IoT, or national Blood Service integration was implemented.",
        "Integrity attribution for direct SQL edits is comparative, not forensic.",
    ])

    heading(doc, "6.3  Recommendations for Future Work", 2)
    add_bullets(doc, [
        "Issue per-facility Ethereum accounts and tighten onlyOwner to a multi-signature or access-control list of hospital wallets.",
        "Deploy the contract to a permissioned or public testnet and retain an explorer link for each hash.",
        "Add an outbox/replay worker so offline anchors are submitted when the node returns.",
        "Integrate with National Blood Service reporting formats and optional BSIS export under Act 1042 [6].",
        "Conduct a formal usability study with laboratory scientists and hospital blood-bank officers.",
        "Optionally add cold-chain IoT sensing for storage and transport, as suggested in related blockchain-IoT literature [38], without placing sensor streams on-chain.",
    ])

    heading(doc, "References", 1)
    refs = [
        '[1]  World Health Organization, "Blood safety and availability," Fact sheet, 12 Jun. 2026. [Online]. Available: https://www.who.int/news-room/fact-sheets/detail/blood-safety-and-availability',
        '[2]  World Health Organization, Global Status Report on Blood Safety and Availability 2025. Geneva, Switzerland: WHO, 2025. [Online]. Available: https://www.who.int/publications/i/item/9789240121546',
        '[3]  Ministry of Health / National Blood Service, Ghana, National Blood Policy, 2nd ed. Accra, Ghana. [Online]. Available: https://nbs.gov.gh/wp-files/NATIONAL%20BLOOD%20POLICY%20(Second%20Edition).pdf',
        '[4]  E. N. Dei and J. K. Ansah, "Blood banking: the situation in Ghana," Transfusion and Apheresis Science, vol. 62, no. 5, Art. no. 103803, Oct. 2023, doi: 10.1016/j.transci.2023.103803.',
        '[5]  AABB, Standards for Blood Banks and Transfusion Services. Bethesda, MD, USA: AABB.',
        '[6]  Republic of Ghana, National Blood Service Act, 2022 (Act 1042).',
        '[7]  Health Facilities Regulatory Agency (HeFRA), "About us." [Online]. Available: https://hefra.gov.gh/about-us/',
        '[8]  Republic of Ghana, Health Institutions and Facilities Act, 2011 (Act 829).',
        '[9]  S. Nakamoto, "Bitcoin: A peer-to-peer electronic cash system," 2008. [Online]. Available: https://bitcoin.org/bitcoin.pdf',
        '[10] Z. Zheng, S. Xie, H. Dai, X. Chen, and H. Wang, "An overview of blockchain technology: Architecture, consensus, and future trends," in Proc. IEEE Int. Congr. Big Data, Honolulu, HI, USA, 2017, pp. 557-564, doi: 10.1109/BigDataCongress.2017.85.',
        '[11] A. Narayanan, J. Bonneau, E. Felten, A. Miller, and S. Goldfeder, Bitcoin and Cryptocurrency Technologies. Princeton, NJ, USA: Princeton Univ. Press, 2016.',
        '[12] G. Wood, "Ethereum: A secure decentralised generalised transaction ledger," Ethereum Project Yellow Paper, 2014. [Online]. Available: https://ethereum.github.io/yellowpaper/paper.pdf',
        '[13] Ethereum Foundation, "Ethereum Virtual Machine (EVM)." [Online]. Available: https://ethereum.org/en/developers/docs/evm/',
        '[14] S. Angraal, J. S. Rumsfeld, and H. M. Krumholz, "Blockchain technology: Applications in health care," Circ. Cardiovasc. Qual. Outcomes, vol. 10, no. 9, p. e003800, 2017, doi: 10.1161/CIRCOUTCOMES.117.003800.',
        '[15] C. C. Agbo, Q. H. Mahmoud, and J. M. Eklund, "Blockchain technology in healthcare: A systematic review," Healthcare, vol. 7, no. 2, p. 56, 2019, doi: 10.3390/healthcare7020056.',
        '[16] D. Stainsby et al., "Serious hazards of transfusion: A decade of hemovigilance in the UK," Transfusion Medicine Reviews, vol. 20, no. 4, pp. 273-282, 2006, doi: 10.1016/j.tmrv.2006.05.002.',
        '[17] M. F. Murphy and J. D. S. Kay, "Patient identification: Problems and potential solutions," Vox Sanguinis, vol. 87, Suppl. 2, pp. S197-S202, 2004, doi: 10.1111/j.1741-6892.2004.00482.x.',
        '[18] A. F. Eder and L. A. Chambers, "Noninfectious complications of blood transfusion," Arch. Pathol. Lab. Med., vol. 131, no. 5, pp. 708-718, 2007.',
        '[19] M. Crosby, P. Nachiappan, P. Pattanayak, S. Verma, and V. Kalyanaraman, "Blockchain technology: Beyond bitcoin," Applied Innovation Review, no. 2, pp. 6-19, 2016.',
        '[20] Z. Zheng, S. Xie, H.-N. Dai, X. Chen, and H. Wang, "Blockchain challenges and opportunities: A survey," Int. J. Web and Grid Services, vol. 14, no. 4, pp. 352-375, 2018.',
        '[21] G. Greenspan, "MultiChain private blockchain - White paper," Coin Sciences Ltd., 2015. [Online]. Available: https://www.multichain.com/download/MultiChain-White-Paper.pdf',
        '[22] N. Szabo, "Smart contracts," 1994. [Online]. Available: https://nakamotoinstitute.org/library/smart-contracts/',
        '[23] N. Szabo, "Smart contracts: Building blocks for digital markets," 1996. [Online]. Available: https://nakamotoinstitute.org/smart-contracts/',
        '[24] W. Viriyasitavat and D. Hoonsopon, "Blockchain characteristics and consensus in modern business processes," Journal of Industrial Information Integration, vol. 13, pp. 32-39, 2019, doi: 10.1016/j.jii.2018.07.004.',
        '[25] X. Yue, H. Wang, D. Jin, M. Li, and W. Jiang, "Healthcare data gateways: Found healthcare intelligence on blockchain with novel privacy risk control," J. Med. Syst., vol. 40, Art. no. 218, 2016, doi: 10.1007/s10916-016-0574-6.',
        '[26] M. Mettler, "Blockchain technology in healthcare: The revolution starts here," in Proc. IEEE Healthcom, 2016, pp. 1-3, doi: 10.1109/HealthCom.2016.7749510.',
        '[27] W. J. Gordon and C. Catalini, "Blockchain technology for healthcare: Facilitating the transition to patient-driven interoperability," Comput. Struct. Biotechnol. J., vol. 16, pp. 224-230, 2018, doi: 10.1016/j.csbj.2018.06.003.',
        '[28] S. Chakraborty, S. Aich, and H.-C. Kim, "A secure healthcare system design framework using blockchain technology," in Proc. 21st Int. Conf. Advanced Communication Technology (ICACT), PyeongChang, Korea (South), 2019, pp. 260-264, doi: 10.23919/ICACT.2019.8701983.',
        '[29] N. Kshetri, "Blockchain\'s roles in meeting key supply chain management objectives," Int. J. Information Management, vol. 39, pp. 80-89, 2018, doi: 10.1016/j.ijinfomgt.2017.12.005.',
        '[30] S. Saberi, M. Kouhizadeh, J. Sarkis, and L. Shen, "Blockchain technology and its potential relationships to sustainable supply chain management," Int. J. Production Research, vol. 57, no. 7, pp. 2117-2135, 2019, doi: 10.1080/00207543.2018.1533261.',
        '[31] S. Shreshtha, S. Rajput, and A. Singh, "Blockchain in blood bank supply management," in Proc. Int. Conf. Innovative Computing & Communication (ICICC), 2021. [Online]. Available: https://ssrn.com/abstract=3879601',
        '[32] N. R. Pradhan, A. P. Singh, and V. Kumar, "Blockchain-enabled traceable, transparent transportation system for blood bank," in Advances in VLSI, Communication, and Signal Processing (Lecture Notes in Electrical Engineering), vol. 683, Springer, 2021, pp. 313-324, doi: 10.1007/978-981-15-6840-4_25.',
        '[33] A. Sanga, W. Mulla, J. Katti, and R. Pise, "Smart blood bank management: Leveraging blockchain for real-time tracking and fraud prevention," in Proc. 9th Int. Conf. Computing, Communication, Control and Automation (ICCUBEA), Pune, India, 2025, pp. 1-6, doi: 10.1109/ICCUBEA65967.2025.11283932.',
        '[34] Laravel LLC, "Laravel documentation," 2026. [Online]. Available: https://laravel.com/docs',
        '[35] Nomic Foundation, "Hardhat documentation." [Online]. Available: https://hardhat.org/docs',
        '[36] ethers.js contributors, "ethers.js documentation," v6. [Online]. Available: https://docs.ethers.org/',
        '[37] Oracle Corporation, "MySQL 8.0 reference manual." [Online]. Available: https://dev.mysql.com/doc/refman/8.0/en/',
        '[38] T. M. Fern\u00e1ndez-Caram\u00e9s and P. Fraga-Lamas, "A review on the use of blockchain for the Internet of Things," IEEE Access, vol. 6, pp. 32979-33001, 2018, doi: 10.1109/ACCESS.2018.2842685.',
    ]
    for ref in refs:
        add_para(doc, ref, first_line=False, space_after=8, align="left")

    heading(doc, "Appendix A  Demo Accounts and Commands", 1)
    add_para(doc, "After php artisan migrate --seed (or migrate:fresh --seed):", first_line=False)
    add_table(
        doc,
        ["Role", "Email", "Password"],
        [
            ["Platform admin", "admin@tarrlok.gh", "TarrlokAdmin2024!"],
            ["Korle Bu hospital", "kwame.mensah@korlebu.gov.gh", "KorleBu2024!"],
            ["Korle Bu laboratory", "ama.osei@korlebu.gov.gh", "KorleBuLab2024!"],
            ["Ridge hospital", "efua.adjei@ridge.gov.gh", "Ridge2024!"],
            ["Ridge laboratory", "kofi.boateng@ridge.gov.gh", "RidgeLab2024!"],
        ],
    )
    add_para(doc, "Public track demo unit after seeding: UNIT-002-00001.", first_line=False)
    add_para(doc, "Blockchain: from blockchain/, run npm run node (keep open) then npm run deploy. In tarrlok/.env set BLOCKCHAIN_ENABLED=true, BLOCKCHAIN_RPC_URL=http://127.0.0.1:8545, and the first Hardhat private key documented in blockchain/README.md. Then php artisan config:clear.", first_line=False)
    add_para(doc, "Useful commands: php artisan test ; php artisan blood:mark-expired ; node scripts/chain-status.js ; node scripts/read-ledger.js.", first_line=False)

    doc.save(OUT)
    print("Wrote", OUT)


if __name__ == "__main__":
    build()
